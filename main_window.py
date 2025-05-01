import os
import json
import sqlite3
import subprocess
import time
import requests
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox, Toplevel, Listbox, MULTIPLE, PhotoImage
#import logging

import customtkinter as ctk
from PIL import Image, ImageTk
from docx import Document

# Modular helper imports
from modules.helpers.window_helper import position_window_at_top
from modules.helpers.template_loader import load_template
from modules.helpers.config_helper import ConfigHelper
from modules.helpers.swarmui_helper import get_available_models
from modules.ui.tooltip import ToolTip
from modules.ui.icon_button import create_icon_button

from modules.generic.generic_list_view import GenericListView
from modules.generic.generic_model_wrapper import GenericModelWrapper
from modules.scenarios.gm_screen_view import GMScreenView
from modules.npcs.npc_graph_editor import NPCGraphEditor
from modules.pcs.pc_graph_editor import PCGraphEditor
from modules.scenarios.scenario_graph_editor import ScenarioGraphEditor
from modules.scenarios.scenario_importer import ScenarioImportWindow
from modules.generic.export_for_foundry import preview_and_export_foundry
from modules.helpers import text_helpers
from db.db import load_schema_from_json, initialize_db
from modules.factions.faction_graph_editor import FactionGraphEditor
from modules.pcs.display_pcs import display_pcs_in_banner
from modules.generic.generic_list_selection_view import GenericListSelectionView


# Set up CustomTkinter appearance
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

# Global process variable for SwarmUI
SWARMUI_PROCESS = None

#logging.basicConfig(level=logging.DEBUG, format='%(levelname)s: %(message)s')


class MainWindow(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("GMCampaignDesigner")
        self.geometry("1920x980")
        self.minsize(1920, 980)
        self.attributes("-fullscreen", True)
        self.current_open_view   = None
        self.current_open_entity = None    # ← initialize here to avoid AttributeError
        initialize_db()
        position_window_at_top(self)
        self.set_window_icon()
        self.create_layout()
        self.load_icons()
        self.create_sidebar()
        self.create_content_area()
        self.create_exit_button()
        self.load_model_config()
        self.init_wrappers()

    # ---------------------------
    # Setup and Layout Methods
    # ---------------------------
    def set_window_icon(self):
        icon_path = os.path.join("assets", "GMCampaignDesigner.ico")
        if os.path.exists(icon_path):
            self.iconbitmap(icon_path)
        icon_image = PhotoImage(file=os.path.join("assets", "GMCampaignDesigner logo.png"))
        self.tk.call('wm', 'iconphoto', self._w, icon_image)

    def create_layout(self):
        self.main_frame = ctk.CTkFrame(self)
        self.main_frame.pack(fill="both", expand=True)

    def load_icons(self):
        self.icons = {
            "change_db": self.load_icon("database_icon.png", size=(48, 48)),
            "swarm_path": self.load_icon("folder_icon.png", size=(48, 48)),
            "manage_scenarios": self.load_icon("scenario_icon.png", size=(48, 48)),
            "manage_pcs": self.load_icon("pc_icon.png", size=(48, 48)),
            "manage_npcs": self.load_icon("npc_icon.png", size=(48, 48)),
            "manage_creatures": self.load_icon("creature_icon.png", size=(48, 48)),
            "manage_factions": self.load_icon("faction_icon.png", size=(48, 48)),
            "manage_places": self.load_icon("places_icon.png", size=(48, 48)),
            "manage_objects": self.load_icon("objects_icon.png", size=(48, 48)),
            "manage_informations": self.load_icon("informations_icon.png", size=(48, 48)),
            "manage_clues": self.load_icon("clues_icon.png", size=(48, 48)),
            "export_scenarios": self.load_icon("export_icon.png", size=(48, 48)),
            "gm_screen": self.load_icon("gm_screen_icon.png", size=(48, 48)),
            "npc_graph": self.load_icon("npc_graph_icon.png", size=(48, 48)),
            "pc_graph": self.load_icon("pc_graph_icon.png", size=(48, 48)),
            "faction_graph": self.load_icon("faction_graph_icon.png", size=(48, 48)),
            "scenario_graph": self.load_icon("scenario_graph_icon.png", size=(48, 48)),
            "generate_portraits": self.load_icon("generate_icon.png", size=(48, 48)),
            "associate_portraits": self.load_icon("associate_icon.png", size=(48, 48)),
            "import_scenario": self.load_icon("import_icon.png", size=(48, 48)),
            "export_foundry": self.load_icon("export_foundry_icon.png", size=(48, 48))
        }

    def load_icon(self, file_name, size=(48, 48)):
        path = os.path.join("assets", file_name)
        try:
            pil_image = Image.open(path)
        except Exception as e:
            #logging.error("Error loading %s: %s", path, e)
            return None
        return ctk.CTkImage(light_image=pil_image, dark_image=pil_image, size=size)

    def create_sidebar(self):
        self.sidebar_frame = ctk.CTkFrame(self.main_frame, width=220)
        self.sidebar_frame.pack(side="left", fill="y", padx=5, pady=5)
        self.sidebar_frame.pack_propagate(False)
        self.sidebar_inner = ctk.CTkFrame(self.sidebar_frame, fg_color="transparent")
        self.sidebar_inner.pack(fill="both", expand=True, padx=5, pady=5)

        # Logo
        logo_path = os.path.join("assets", "GMCampaignDesigner logo.png")
        if os.path.exists(logo_path):
            logo_image = Image.open(logo_path).resize((60, 60))
            logo = ctk.CTkImage(light_image=logo_image, dark_image=logo_image, size=(100, 100))
            self.logo_image = logo
            logo_label = ctk.CTkLabel(self.sidebar_inner, image=logo, text="")
            logo_label.pack(pady=(0, 3), anchor="center")

        # Header label
        header_label = ctk.CTkLabel(self.sidebar_inner, text="Campaign Tools", font=("Helvetica", 16, "bold"))
        header_label.pack(pady=(0, 5), anchor="center")

        # Database display container
        db_container = ctk.CTkFrame(self.sidebar_inner, fg_color="transparent",
                                    border_color="#005fa3", border_width=2, corner_radius=8)
        db_container.pack(pady=(0, 5), anchor="center", fill="x", padx=5)
        db_title_label = ctk.CTkLabel(db_container, text="Database:", font=("Segoe UI", 16, "bold"),
                                    fg_color="transparent", text_color="white")
        db_title_label.pack(pady=(3, 0), anchor="center")
        db_path = ConfigHelper.get("Database", "path", fallback="default_campaign.db")
        db_name = os.path.splitext(os.path.basename(db_path))[0]
        self.db_name_label = ctk.CTkLabel(db_container, text=db_name,
                                        font=("Segoe UI", 14, "italic"), fg_color="transparent", text_color="white")
        self.db_name_label.pack(pady=(0, 3), anchor="center")

        self.create_icon_grid()

    def create_icon_grid(self):
        icons_frame = ctk.CTkFrame(self.sidebar_inner, fg_color="transparent")
        icons_frame.pack(fill="both", expand=True, padx=5, pady=5)
        columns = 2
        for col in range(columns):
            icons_frame.grid_columnconfigure(col, weight=1)
        icons_list = [
            ("change_db", "Change Data Storage", self.change_database_storage),
            ("swarm_path", "Set SwarmUI Path", self.select_swarmui_path),
            ("manage_scenarios", "Manage Scenarios", lambda: self.open_entity("scenarios")),
            ("manage_pcs", "Manage PCs", lambda: self.open_entity("pcs")),
            ("manage_npcs", "Manage NPCs", lambda: self.open_entity("npcs")),
            ("manage_creatures", "Manage Creatures", lambda: self.open_entity("creatures")),
            ("manage_factions", "Manage Factions", lambda: self.open_entity("factions")),
            ("manage_places", "Manage Places", lambda: self.open_entity("places")),
            ("manage_objects", "Manage Objects", lambda: self.open_entity("objects")),
            ("manage_informations", "Manage Informations", lambda: self.open_entity("informations")),
            ("manage_clues", "Manage Clues", lambda: self.open_entity("clues")),
            ("export_scenarios", "Export Scenarios", self.preview_and_export_scenarios),
            ("gm_screen", "Open GM Screen", self.open_gm_screen),
            ("npc_graph", "Open NPC Graph Editor", self.open_npc_graph_editor),
            ("pc_graph", "Open PC Graph Editor", self.open_pc_graph_editor),
            ("faction_graph", "Open Factions Graph Editor", self.open_faction_graph_editor),
            ("scenario_graph", "Open Scenario Graph Editor", self.open_scenario_graph_editor),
            ("generate_portraits", "Generate Portraits", self.generate_missing_portraits),
            ("associate_portraits", "Associate NPC Portraits", self.associate_npc_portraits),
            ("import_scenario", "Import Scenario", self.open_scenario_importer),
            ("export_foundry", "Export Scenarios for Foundry", self.export_foundry)
        ]
        self.icon_buttons = []
        for idx, (icon_key, tooltip, cmd) in enumerate(icons_list):
            row = idx // columns
            col = idx % columns
            btn = create_icon_button(icons_frame, self.icons[icon_key], tooltip, cmd)
            btn.grid(row=row, column=col, padx=5, pady=5)
            self.icon_buttons.append(btn)

    def create_content_area(self):
        self.content_frame = ctk.CTkFrame(self.main_frame)
        self.content_frame.pack(side="right", fill="both", expand=True, padx=10, pady=10)

        self.content_frame.grid_rowconfigure(0, weight=1)
        self.content_frame.grid_rowconfigure(1, weight=0)
        self.content_frame.grid_columnconfigure(0, weight=1)

        # ✅ Always explicitly create these initially:
        self.banner_frame = ctk.CTkFrame(self.content_frame, height=150, fg_color="#444")
        self.inner_content_frame = ctk.CTkFrame(self.content_frame, fg_color="#222")

        self.banner_toggle_btn = ctk.CTkButton(
            self.sidebar_inner,
            text="▼",
            width=40,
            height=30,
            command=self._toggle_banner,
            fg_color="#555",
            hover_color="#777",
            font=("", 16)
        )
        self.banner_toggle_btn.place(relx=1.0, rely=0.0, anchor="ne")

        self.banner_visible = False
        self.current_open_view = None

    def _toggle_banner(self):
        # --- GRAPH MODE?  (no current_open_entity but a graph_type set) ---
        if self.current_open_entity is None and getattr(self, "_graph_type", None):
            # snapshot existing graph state
            old_container = self.current_open_view
            old_editor    = getattr(old_container, "graph_editor", None)
            state         = old_editor.get_state() if old_editor else None

            # hide or show banner + inner frames
            if self.banner_visible:
                self.banner_frame.grid_remove()
                self.inner_content_frame.grid_remove()
                self.content_frame.grid_rowconfigure(0, weight=1)
                self.content_frame.grid_rowconfigure(1, weight=0)
                self.banner_visible = False
                self.banner_toggle_btn.configure(text="▼")
            else:
                self.banner_frame.grid(row=0, column=0, sticky="ew")
                self.inner_content_frame.grid(row=1, column=0, sticky="nsew")
                display_pcs_in_banner(
                    self.banner_frame,
                    {pc["Name"]: pc for pc in self.pc_wrapper.load_items()}
                )
                self.inner_content_frame.grid_rowconfigure(0, weight=1)
                self.inner_content_frame.grid_columnconfigure(0, weight=1)
                self.content_frame.grid_rowconfigure(0, weight=0)
                self.content_frame.grid_rowconfigure(1, weight=1)
                self.banner_visible = True
                self.banner_toggle_btn.configure(text="▲")

            # destroy the old container (it was still parented in the wrong frame)
            old_container.destroy()

            # now re‐create the same graph under the correct container
            parent = self.get_content_container()
            new_container = ctk.CTkFrame(parent)
            new_container.grid(row=0, column=0, sticky="nsew")
            parent.grid_rowconfigure(0, weight=1)
            parent.grid_columnconfigure(0, weight=1)

            # re‐instantiate the proper editor type, then restore its state
            if self._graph_type == 'npc':
                editor = NPCGraphEditor(new_container, self.npc_wrapper, self.faction_wrapper)
            elif self._graph_type == 'pc':
                editor = PCGraphEditor(new_container, self.pc_wrapper, self.faction_wrapper)
            elif self._graph_type == 'faction':
                editor = FactionGraphEditor(new_container, self.faction_wrapper)
            else:  # 'scenario'
                editor = ScenarioGraphEditor(
                    new_container,
                    GenericModelWrapper("scenarios"),
                    GenericModelWrapper("npcs"),
                    GenericModelWrapper("creatures"),
                    GenericModelWrapper("places")
                )

            editor.pack(fill="both", expand=True)
            if state is not None and hasattr(editor, "set_state"):
                editor.set_state(state)

            # save the new container/editor
            new_container.graph_editor = editor
            self.current_open_view     = new_container
            # leave current_open_entity = None

            return  # end of graph‐mode toggle
        if self.banner_visible:
            # COLLAPSE BANNER
            if self.banner_frame.winfo_exists():
                self.banner_frame.grid_remove()
            if self.inner_content_frame.winfo_exists():
                self.inner_content_frame.grid_remove()

            if self.current_open_view:
                # Save entity and then destroy
                entity = self.current_open_entity
                self.current_open_view.destroy()

                # Re-create content in content_frame
                self.current_open_view = ctk.CTkFrame(self.content_frame)
                self.current_open_view.grid(row=0, column=0, sticky="nsew")

                wrapper = GenericModelWrapper(entity)
                template = load_template(entity)
                view = GenericListView(self.current_open_view, wrapper, template)
                view.pack(fill="both", expand=True)

                load_button = ctk.CTkButton(
                    self.current_open_view,
                    text=f"Load {entity.capitalize()}",
                    command=lambda: self.load_items_from_json(view, entity)
                )
                load_button.pack(side="right", padx=(5,5))
                # Assuming `editor_window` is your CTkToplevel or CTkFrame
                save_btn = ctk.CTkButton(
                    self.current_open_view,
                    text=f"Save {entity.capitalize()}",
                    command=lambda: self.save_items_to_json(view, entity)
                )
                save_btn.pack(side="right", padx=(5,5))

            self.content_frame.grid_rowconfigure(0, weight=1)
            self.content_frame.grid_rowconfigure(1, weight=0)

            self.banner_visible = False
            self.banner_toggle_btn.configure(text="▼")
        else:
            # EXPAND BANNER
            if not self.banner_frame.winfo_exists():
                self.banner_frame = ctk.CTkFrame(self.content_frame, height=150, fg_color="#444")

            if not self.inner_content_frame.winfo_exists():
                self.inner_content_frame = ctk.CTkFrame(self.content_frame, fg_color="#222")

            self.banner_frame.grid(row=0, column=0, sticky="ew")
            self.inner_content_frame.grid(row=1, column=0, sticky="nsew")
            pcs_items = {pc["Name"]: pc for pc in self.pc_wrapper.load_items()}
            display_pcs_in_banner(self.banner_frame, pcs_items)

            # ✅ CRITICAL FIX: make inner_content_frame fully expandable
            self.inner_content_frame.grid_rowconfigure(0, weight=1)
            self.inner_content_frame.grid_columnconfigure(0, weight=1)

            if self.current_open_view:
                entity = self.current_open_entity
                self.current_open_view.destroy()

                self.current_open_view = ctk.CTkFrame(self.inner_content_frame)
                self.current_open_view.grid(row=0, column=0, sticky="nsew")

                wrapper = GenericModelWrapper(entity)
                template = load_template(entity)
                view = GenericListView(self.current_open_view, wrapper, template)
                view.pack(fill="both", expand=True)

                load_button = ctk.CTkButton(
                    self.current_open_view,
                    text=f"Load {entity.capitalize()}",
                    command=lambda: self.load_items_from_json(view, entity)
                )
                load_button.pack(side="right", padx=(5,5))
                # Assuming `editor_window` is your CTkToplevel or CTkFrame
                save_btn = ctk.CTkButton(
                    self.current_open_view,
                    text=f"Save {entity.capitalize()}",
                    command=lambda: self.save_items_to_json(view, entity)
                )
                save_btn.pack(side="right", padx=(5,5))

            self.content_frame.grid_rowconfigure(0, weight=0)
            self.content_frame.grid_rowconfigure(1, weight=1)

            self.banner_visible = True
            self.banner_toggle_btn.configure(text="▲")

    def get_content_container(self):
        """Choose correct parent depending on banner state."""
        if self.banner_visible:
            return self.inner_content_frame
        else:
            return self.content_frame
    def create_exit_button(self):
        exit_button = ctk.CTkButton(self, text="✕", command=self.destroy,
                                    fg_color="red", hover_color="#AA0000",
                                    width=20, height=20, corner_radius=15)
        exit_button.place(relx=0.9999, rely=0.01, anchor="ne")

    def load_model_config(self):
        self.models_path = ConfigHelper.get("Paths", "models_path",
                                            fallback=r"E:\SwarmUI\SwarmUI\Models\Stable-diffusion")
        self.model_options = get_available_models()

    def init_wrappers(self):
        self.place_wrapper = GenericModelWrapper("places")
        self.npc_wrapper = GenericModelWrapper("npcs")
        self.pc_wrapper = GenericModelWrapper("pcs")
        self.faction_wrapper = GenericModelWrapper("factions")
        self.object_wrapper = GenericModelWrapper("objects")
        self.creature_wrapper = GenericModelWrapper("creatures")

    def open_faction_graph_editor(self):
        self._graph_type = 'faction'
        self.clear_current_content()
        self.banner_toggle_btn.configure(state="normal")
        parent = self.get_content_container()

        container = ctk.CTkFrame(parent)
        container.grid(row=0, column=0, sticky="nsew")
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)

        editor = FactionGraphEditor(container, self.faction_wrapper)
        editor.pack(fill="both", expand=True)

        # keep both container and editor so we can snapshot/restore
        container.graph_editor = editor
        self.current_open_view   = container
        self.current_open_entity = None
        

    # =============================================================
    # Methods Called by Icon Buttons (Event Handlers)
    # =============================================================
    def clear_current_content(self):
        if self.banner_visible:
            # If banner is visible, clear only inner_content_frame
            for widget in self.inner_content_frame.winfo_children():
                widget.destroy()
        else:
            # If banner is hidden, clear ONLY dynamic children of content_frame, NOT banner_frame
            for widget in self.content_frame.winfo_children():
                if widget not in (self.banner_frame, self.inner_content_frame):
                    widget.destroy()

    def move_current_view(self):
        """Move the current open view to the correct container based on banner state."""
        if self.current_open_view is not None:
            try:
                self.current_open_view.grid_forget()
            except tk.TclError:
                # If the widget was destroyed (because it was inside inner_content_frame), clear it
                self.current_open_view = None
                return

            parent = self.get_content_container()
            self.current_open_view.master = parent
            self.current_open_view.grid(row=0, column=0, sticky="nsew")

    def open_entity(self, entity):
        self.clear_current_content()
        target_parent = self.get_content_container()
        self.banner_toggle_btn._state="normal"
        container = ctk.CTkFrame(target_parent)
        container.grid(row=0, column=0, sticky="nsew")
        self.current_open_view = container
        self.current_open_entity = entity  # ✅ Add this clearly!

        wrapper = GenericModelWrapper(entity)
        template = load_template(entity)
        view = GenericListView(container, wrapper, template)
        view.pack(fill="both", expand=True)

        load_button = ctk.CTkButton(
            container,
            text=f"Load {entity.capitalize()}",
            command=lambda: self.load_items_from_json(view, entity)
        )
        load_button.pack(side="right", padx=(5,5))
        # Assuming `editor_window` is your CTkToplevel or CTkFrame
        save_btn = ctk.CTkButton(
            container,
            text=f"Save {entity.capitalize()}",
            command=lambda: self.save_items_to_json(view, entity)
        )
        save_btn.pack(side="right", padx=(5,5))
    
    def save_items_to_json(self, view, entity_name):
        # 1) Ask the user where to save
        path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")],
            title=f"Export {entity_name.capitalize()} to JSON"
        )
        if not path:
            return  # user hit “Cancel”

        # 2) Grab the items from the view if possible…
        try:
            # GenericListView *might* have a method or attribute that holds its current items
            items = view.get_items()                # ← if you’ve added a get_items()
        except AttributeError:
            try:
                items = view.items                  # ← or maybe it’s stored in view.items
            except Exception:
                # 3) …otherwise fall back on the DB
                wrapper = GenericModelWrapper(entity_name)
                items   = wrapper.load_items()

        # 4) Serialize to JSON
        data = { entity_name: items }
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save {entity_name}:\n{e}")
            return

        # 5) Let the user know it worked
        messagebox.showinfo("Export Successful", f"Wrote {len(items)} {entity_name} to:\n{path}")

    def load_items_from_json(self, view, entity_name):
        file_path = filedialog.askopenfilename(
            title=f"Load {entity_name.capitalize()} from JSON",
            filetypes=[("JSON Files", "*.json"), ("All Files", "*.*")]
        )
        if not file_path:
            return
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                data = json.load(file)
                items = data.get(entity_name, [])
                view.add_items(items)
                messagebox.showinfo("Success", f"{len(items)} {entity_name} loaded successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load {entity_name}: {e}")

    def open_gm_screen(self):
        # 1) Clear any existing content
        self.clear_current_content()
        # 2) Load all scenarios
        scenario_wrapper = GenericModelWrapper("scenarios")
        scenarios = scenario_wrapper.load_items()
        if not scenarios:
            messagebox.showwarning("No Scenarios", "No scenarios available.")
            return

        # 3) Ensure the PC‐banner is shown and up to date
        if not self.banner_frame.winfo_ismapped():
            self.banner_frame.grid(row=0, column=0, sticky="ew")
        pcs_items = {pc["Name"]: pc for pc in self.pc_wrapper.load_items()}
        if pcs_items:
            display_pcs_in_banner(self.banner_frame, pcs_items)

        # 4) Prepare inner content area
        self.inner_content_frame.grid(row=1, column=0, sticky="nsew")
        for w in self.inner_content_frame.winfo_children():
            w.destroy()
        parent = self.inner_content_frame

        # 5) Callback to open a selected scenario in detail
        def on_scenario_select(entity_type, entity_name):
            selected = next(
                (s for s in scenarios
                if s.get("Name", s.get("Title", "")) == entity_name),
                None
            )
            if not selected:
                messagebox.showwarning("Not Found", f"Scenario '{entity_name}' not found.")
                return
            # clear list and show scenario detail
            for w in parent.winfo_children():
                w.destroy()
            detail_container = ctk.CTkFrame(parent)
            detail_container.grid(row=0, column=0, sticky="nsew")
            view = GMScreenView(detail_container, scenario_item=selected)
            view.pack(fill="both", expand=True)

        # 6) Insert the generic list‐selection view
        list_selection = GenericListSelectionView(
            parent,
            "scenarios",
            scenario_wrapper,
            load_template("scenarios"),
            on_select_callback=on_scenario_select
        )
        list_selection.pack(fill="both", expand=True)

        # 7) Lock banner and configure grid weights
        self.banner_visible = True
        self.banner_toggle_btn.configure(text="▲")
        self.banner_toggle_btn._state = "disabled"

        # Make row 0 (banner) fixed height, row 1 (content) expand
        self.content_frame.grid_rowconfigure(0, weight=0)
        self.content_frame.grid_rowconfigure(1, weight=1)
        self.content_frame.grid_columnconfigure(0, weight=1)

        # Make the inner_content_frame fully fill its cell
        self.inner_content_frame.grid_rowconfigure(0, weight=1)
        self.inner_content_frame.grid_columnconfigure(0, weight=1)

    def open_npc_graph_editor(self):
        self._graph_type = 'npc'
        self.clear_current_content()
        self.banner_toggle_btn.configure(state="normal")
        parent = self.get_content_container()

        container = ctk.CTkFrame(parent)
        container.grid(row=0, column=0, sticky="nsew")
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)

        editor = NPCGraphEditor(container, self.npc_wrapper, self.faction_wrapper)
        editor.pack(fill="both", expand=True)

        # keep both container and editor so we can snapshot/restore
        container.graph_editor = editor
        self.current_open_view   = container
        self.current_open_entity = None


    def open_pc_graph_editor(self):
        self._graph_type = 'pc'
        self.clear_current_content()
        self.banner_toggle_btn.configure(state="normal")
        parent = self.get_content_container()

        container = ctk.CTkFrame(parent)
        container.grid(row=0, column=0, sticky="nsew")
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)

        editor = PCGraphEditor(container, self.pc_wrapper, self.faction_wrapper)
        editor.pack(fill="both", expand=True)

        container.graph_editor = editor
        self.current_open_view   = container
        self.current_open_entity = None


    def open_scenario_graph_editor(self):
        self._graph_type = 'scenario'
        self.clear_current_content()
        self.banner_toggle_btn.configure(state="normal")
        parent = self.get_content_container()

        container = ctk.CTkFrame(parent)
        container.grid(row=0, column=0, sticky="nsew")
        parent.grid_rowconfigure(0, weight=1)
        parent.grid_columnconfigure(0, weight=1)

        editor = ScenarioGraphEditor(
            container,
            GenericModelWrapper("scenarios"),
            GenericModelWrapper("npcs"),
            GenericModelWrapper("creatures"),
            GenericModelWrapper("places")
        )
        editor.pack(fill="both", expand=True)

        container.graph_editor = editor
        self.current_open_view   = container
        self.current_open_entity = None

    def export_foundry(self):
        preview_and_export_foundry(self)

    def open_scenario_importer(self):
        self.clear_current_content()
        container = ctk.CTkFrame(self.content_frame)
        container.grid(row=0, column=0, sticky="nsew")
        ScenarioImportWindow(container)

    def change_database_storage(self):
        # 1) Pick or create .db
        choice = messagebox.askquestion(
            "Change Database",
            "Do you want to open an existing database file?"
        )
        if choice == "yes":
            new_db_path = filedialog.askopenfilename(
                title="Select Database",
                filetypes=[("SQLite DB Files", "*.db"), ("All Files", "*.*")]
            )
        else:
            new_db_path = filedialog.asksaveasfilename(
                title="Create New Database",
                defaultextension=".db",
                filetypes=[("SQLite DB Files", "*.db"), ("All Files", "*.*")]
            )
        if not new_db_path:
            return

        # 2) Persist to config so get_connection()/init_db() will pick it up
        ConfigHelper.set("Database", "path", new_db_path)

        # 3) Open a fresh connection and create all tables based on JSON templates
        conn = sqlite3.connect(new_db_path)
        cursor = conn.cursor()

        # For each entity, load its template and build a CREATE TABLE
        for entity in ("pcs","npcs", "scenarios", "factions",
                    "places", "objects", "creatures", "informations","clues"):

            tpl = load_template(entity)   # loads modules/<entity>/<entity>_template.json
            cols = []
            for i, field in enumerate(tpl["fields"]):
                name = field["name"]
                ftype = field["type"]
                # map JSON -> SQL
                if ftype in ("text", "longtext"):
                    sql_type = "TEXT"
                elif ftype == "boolean":
                    sql_type = "BOOLEAN"
                elif ftype == "list":
                    # we store lists as JSON strings
                    sql_type = "TEXT"
                elif ftype == "file":
                    # we store lists as JSON strings
                    sql_type = "TEXT"
                else:
                    sql_type = "TEXT"

                # first field is primary key
                if i == 0:
                    cols.append(f"{name} {sql_type} PRIMARY KEY")
                else:
                    cols.append(f"{name} {sql_type}")

            ddl = f"CREATE TABLE IF NOT EXISTS {entity} ({', '.join(cols)})"
            cursor.execute(ddl)

        # 4) Re‑create the graph viewer tables
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS nodes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                npc_name TEXT,
                x INTEGER,
                y INTEGER,
                color TEXT
            )
        """)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS links (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                npc_name1 TEXT,
                npc_name2 TEXT,
                text TEXT,
                arrow_mode TEXT
            )
        """)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS shapes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                type TEXT,
                x INTEGER,
                y INTEGER,
                w INTEGER,
                h INTEGER,
                color TEXT,
                tag TEXT,
                z INTEGER
            )
        """)

        conn.commit()
        conn.close()

        # 5) Re‑initialise your in‑memory wrappers & update the label
        #    (and run any schema‐migrations if you still need them)
        self.place_wrapper    = GenericModelWrapper("places")
        self.npc_wrapper      = GenericModelWrapper("npcs")
        self.pc_wrapper      = GenericModelWrapper("pcs")
        self.faction_wrapper  = GenericModelWrapper("factions")
        self.object_wrapper   = GenericModelWrapper("objects")
        self.creature_wrapper = GenericModelWrapper("creatures")
        self.information_wrapper = GenericModelWrapper("informations")
        self.clues_wrapper = GenericModelWrapper("clues")

        db_name = os.path.splitext(os.path.basename(new_db_path))[0]
        self.db_name_label.configure(text=db_name)

    def select_swarmui_path(self):
        folder = filedialog.askdirectory(title="Select SwarmUI Path")
        if folder:
            ConfigHelper.set("Paths", "swarmui_path", folder)
            messagebox.showinfo("SwarmUI Path Set", f"SwarmUI path set to:\n{folder}")

    def launch_swarmui(self):
        global SWARMUI_PROCESS
        swarmui_path = ConfigHelper.get("Paths", "swarmui_path", fallback=r"E:\SwarmUI\SwarmUI")
        SWARMUI_CMD = os.path.join(swarmui_path, "launch-windows.bat")
        env = os.environ.copy()
        env.pop('VIRTUAL_ENV', None)
        if SWARMUI_PROCESS is None or SWARMUI_PROCESS.poll() is not None:
            try:
                SWARMUI_PROCESS = subprocess.Popen(
                    SWARMUI_CMD,
                    shell=True,
                    cwd=swarmui_path,
                    env=env
                )
                time.sleep(120.0)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to launch SwarmUI: {e}")

    def cleanup_swarmui(self):
        global SWARMUI_PROCESS
        if SWARMUI_PROCESS is not None and SWARMUI_PROCESS.poll() is None:
            SWARMUI_PROCESS.terminate()

    # ------------------------------------------------------
    # Unified Generate Portraits for NPCs and Creatures
    # ------------------------------------------------------
    def generate_missing_portraits(self):
        top = ctk.CTkToplevel(self)
        top.title("Generate Portraits")
        top.geometry("300x150")
        top.transient(self)
        top.grab_set()
        # Use ctk.StringVar (not CTkStringVar)
        selection = ctk.StringVar(value="NPC")
        ctk.CTkLabel(top, text="Generate portraits for:").pack(pady=10)
        ctk.CTkRadioButton(top, text="NPCs", variable=selection, value="NPC").pack(pady=5)
        ctk.CTkRadioButton(top, text="Creatures", variable=selection, value="Creature").pack(pady=5)
        def on_confirm():
            choice = selection.get()
            top.destroy()
            if choice == "NPC":
                self.generate_missing_npc_portraits()
            else:
                self.generate_missing_creature_portraits()
        ctk.CTkButton(top, text="Continue", command=on_confirm).pack(pady=10)

    def generate_missing_npc_portraits(self):
        def confirm_model_and_continue():
            ConfigHelper.set("LastUsed", "model", self.selected_model.get())
            top.destroy()
            self.generate_portraits_continue_npcs()
        top = ctk.CTkToplevel(self)
        top.title("Select AI Model for NPCs")
        top.geometry("400x200")
        top.transient(self)
        top.grab_set()
        ctk.CTkLabel(top, text="Select AI Model to use for NPC portrait generation:").pack(pady=20)
        last_model = ConfigHelper.get("LastUsed", "model", fallback=None)
        # Use ctk.StringVar
        if last_model in self.model_options:
            self.selected_model = ctk.StringVar(value=last_model)
        else:
            self.selected_model = ctk.StringVar(value=self.model_options[0])
        ctk.CTkOptionMenu(top, values=self.model_options, variable=self.selected_model).pack(pady=10)
        ctk.CTkButton(top, text="Continue", command=confirm_model_and_continue).pack(pady=10)

    def generate_portraits_continue_npcs(self):
        db_path = ConfigHelper.get("Database", "path", fallback="default_campaign.db")
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM npcs")
        npc_rows = cursor.fetchall()
        modified = False
        for npc in npc_rows:
            portrait = npc["Portrait"] if npc["Portrait"] is not None else ""
            if not portrait.strip():
                npc_dict = dict(npc)
                self.generate_portrait_for_npc(npc_dict)
                if npc_dict.get("Portrait"):
                    cursor.execute("UPDATE npcs SET Portrait = ? WHERE Name = ?", (npc_dict["Portrait"], npc["Name"]))
                    modified = True
        if modified:
            conn.commit()
            print("Updated NPC database with generated portraits.")
        else:
            print("No NPCs were missing portraits.")
        conn.close()

    def generate_missing_creature_portraits(self):
        def confirm_model_and_continue():
            ConfigHelper.set("LastUsed", "model", self.selected_model.get())
            top.destroy()
            self.generate_portraits_continue_creatures()
        top = ctk.CTkToplevel(self)
        top.title("Select AI Model for Creatures")
        top.geometry("400x200")
        top.transient(self)
        top.grab_set()
        ctk.CTkLabel(top, text="Select AI Model to use for creature portrait generation:").pack(pady=20)
        last_model = ConfigHelper.get("LastUsed", "model", fallback=None)
        if last_model in self.model_options:
            self.selected_model = ctk.StringVar(value=last_model)
        else:
            self.selected_model = ctk.StringVar(value=self.model_options[0])
        ctk.CTkOptionMenu(top, values=self.model_options, variable=self.selected_model).pack(pady=10)
        ctk.CTkButton(top, text="Continue", command=confirm_model_and_continue).pack(pady=10)

    def generate_portraits_continue_creatures(self):
        db_path = ConfigHelper.get("Database", "path", fallback="default_campaign.db")
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM creatures")
        creature_rows = cursor.fetchall()
        modified = False
        for creature in creature_rows:
            portrait = creature["Portrait"] if creature["Portrait"] is not None else ""
            if not portrait.strip():
                creature_dict = dict(creature)
                self.generate_portrait_for_creature(creature_dict)
                if creature_dict.get("Portrait"):
                    cursor.execute("UPDATE creatures SET Portrait = ? WHERE Name = ?", (creature_dict["Portrait"], creature_dict["Name"]))
                    modified = True
        if modified:
            conn.commit()
            print("Updated creature database with generated portraits.")
        else:
            print("No creatures were missing portraits.")
        conn.close()

    def generate_portrait_for_npc(self, npc):
        self.launch_swarmui()
        SWARM_API_URL = "http://127.0.0.1:7801"
        try:
            session_url = f"{SWARM_API_URL}/API/GetNewSession"
            session_response = requests.post(session_url, json={}, headers={"Content-Type": "application/json"})
            session_data = session_response.json()
            session_id = session_data.get("session_id")
            if not session_id:
                print(f"Failed to obtain session ID for NPC {npc.get('Name', 'Unknown')}")
                return
            npc_name = npc.get("Name", "Unknown")
            npc_role = npc.get("Role", "Unknown")
            npc_faction = npc.get("Factions", "Unknown")
            npc_desc = npc.get("Description", "Unknown")
            npc_desc = text_helpers.format_longtext(npc_desc)
            prompt = f"{npc_name} {npc_desc} {npc_role} {npc_faction}"
            prompt_data = {
                "session_id": session_id,
                "images": 1,
                "prompt": prompt,
                "negativeprompt": ("blurry, low quality, comics style, mangastyle, paint style, watermark, ugly, "
                                "monstrous, too many fingers, too many legs, too many arms, bad hands, "
                                "unrealistic weapons, bad grip on equipment, nude"),
                "model": self.selected_model.get(),
                "width": 1024,
                "height": 1024,
                "cfgscale": 9,
                "steps": 20,
                "seed": -1
            }
            generate_url = f"{SWARM_API_URL}/API/GenerateText2Image"
            image_response = requests.post(generate_url, json=prompt_data, headers={"Content-Type": "application/json"})
            image_data = image_response.json()
            images = image_data.get("images")
            if not images or len(images) == 0:
                print(f"Image generation failed for NPC '{npc_name}'")
                return
            image_url = f"{SWARM_API_URL}/{images[0]}"
            downloaded_image = requests.get(image_url)
            if downloaded_image.status_code != 200:
                print(f"Failed to download generated image for NPC '{npc_name}'")
                return
            output_filename = f"{npc_name.replace(' ', '_')}_portrait.png"
            with open(output_filename, "wb") as f:
                f.write(downloaded_image.content)
            GENERATED_FOLDER = "assets/generated"
            os.makedirs(GENERATED_FOLDER, exist_ok=True)
            shutil.copy(output_filename, os.path.join(GENERATED_FOLDER, output_filename))
            npc["Portrait"] = self.copy_and_resize_portrait(npc, output_filename)
            os.remove(output_filename)
            print(f"Generated portrait for NPC '{npc_name}'")
        except Exception as e:
            print(f"Error generating portrait for NPC '{npc.get('Name', 'Unknown')}': {e}")

    def generate_portrait_for_creature(self, creature):
        self.launch_swarmui()
        SWARM_API_URL = "http://127.0.0.1:7801"
        try:
            session_url = f"{SWARM_API_URL}/API/GetNewSession"
            session_response = requests.post(session_url, json={}, headers={"Content-Type": "application/json"})
            session_data = session_response.json()
            session_id = session_data.get("session_id")
            if not session_id:
                print(f"Failed to obtain session ID for Creature {creature.get('Name', 'Unknown')}")
                return
            creature_name = creature.get("Name", "Unknown")
            creature_desc = creature.get("Description", "Unknown")
            stats = creature.get("Stats", "")
            creature_desc_formatted = text_helpers.format_longtext(creature_desc)
            prompt = f"{creature_name} {creature_desc_formatted} {stats}"
            prompt_data = {
                "session_id": session_id,
                "images": 1,
                "prompt": prompt,
                "negativeprompt": ("blurry, low quality, comics style, mangastyle, paint style, watermark, ugly, "
                                "monstrous, too many fingers, too many legs, too many arms, bad hands, "
                                "unrealistic weapons, bad grip on equipment, nude"),
                "model": self.selected_model.get(),
                "width": 1024,
                "height": 1024,
                "cfgscale": 9,
                "steps": 20,
                "seed": -1
            }
            generate_url = f"{SWARM_API_URL}/API/GenerateText2Image"
            image_response = requests.post(generate_url, json=prompt_data, headers={"Content-Type": "application/json"})
            image_data = image_response.json()
            images = image_data.get("images")
            if not images or len(images) == 0:
                print(f"Image generation failed for Creature '{creature_name}'")
                return
            image_url = f"{SWARM_API_URL}/{images[0]}"
            downloaded_image = requests.get(image_url)
            if downloaded_image.status_code != 200:
                print(f"Failed to download generated image for Creature '{creature_name}'")
                return
            output_filename = f"{creature_name.replace(' ', '_')}_portrait.png"
            with open(output_filename, "wb") as f:
                f.write(downloaded_image.content)
            GENERATED_FOLDER = "assets/generated"
            os.makedirs(GENERATED_FOLDER, exist_ok=True)
            shutil.copy(output_filename, os.path.join(GENERATED_FOLDER, output_filename))
            creature["Portrait"] = self.copy_and_resize_portrait(creature, output_filename)
            os.remove(output_filename)
            print(f"Generated portrait for Creature '{creature_name}'")
        except Exception as e:
            print(f"Error generating portrait for Creature '{creature.get('Name', 'Unknown')}': {e}")

    def copy_and_resize_portrait(self, entity, src_path):
        PORTRAIT_FOLDER = "assets/portraits"
        MAX_PORTRAIT_SIZE = (1024, 1024)
        os.makedirs(PORTRAIT_FOLDER, exist_ok=True)
        name = entity.get("Name", "Unnamed").replace(" ", "_")
        ext = os.path.splitext(src_path)[-1].lower()
        dest_filename = f"{name}_{id(self)}{ext}"
        dest_path = os.path.join(PORTRAIT_FOLDER, dest_filename)
        with Image.open(src_path) as img:
            img = img.convert("RGB")
            img.thumbnail(MAX_PORTRAIT_SIZE)
            img.save(dest_path)
        return dest_path

    def preview_and_export_scenarios(self):
        scenario_wrapper = GenericModelWrapper("scenarios")
        scenario_items = scenario_wrapper.load_items()
        if not scenario_items:
            messagebox.showwarning("No Scenarios", "There are no scenarios available.")
            return
        selection_window = Toplevel()
        selection_window.title("Select Scenarios to Export")
        selection_window.geometry("400x300")
        listbox = Listbox(selection_window, selectmode="multiple", height=15)
        listbox.pack(fill="both", expand=True, padx=10, pady=10)
        scenario_titles = [scenario.get("Title", "Unnamed Scenario") for scenario in scenario_items]
        for title in scenario_titles:
            listbox.insert("end", title)
        def export_selected():
            selected_indices = listbox.curselection()
            if not selected_indices:
                messagebox.showwarning("No Selection", "Please select at least one scenario to export.")
                return
            selected_scenarios = [scenario_items[i] for i in selected_indices]
            self.preview_and_save(selected_scenarios)
            selection_window.destroy()
        ctk.CTkButton(selection_window, text="Export Selected", command=export_selected).pack(pady=5)

    def preview_and_save(self, selected_scenarios):
        creature_items = {creature["Name"]: creature for creature in self.creature_wrapper.load_items()}
        place_items = {place["Name"]: place for place in self.place_wrapper.load_items()}
        npc_items = {npc["Name"]: npc for npc in self.npc_wrapper.load_items()}

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Files", "*.docx"), ("All Files", "*.*")],
            title="Save Scenario Export"
        )
        if not file_path:
            return

        doc = Document()
        doc.add_heading("Campaign Scenarios", level=1)
        for scenario in selected_scenarios:
            title = scenario.get("Title", "Unnamed Scenario")
            summary = scenario.get("Summary", "No description provided.")
            secrets = scenario.get("Secrets", "No secrets provided.")

            doc.add_heading(title, level=2)
            doc.add_heading("Summary", level=3)
            if isinstance(summary, dict):
                p = doc.add_paragraph()
                run = p.add_run(summary.get("text", ""))
                self.apply_formatting(run, summary.get("formatting", {}))
            else:
                doc.add_paragraph(str(summary))

            doc.add_heading("Secrets", level=3)
            if isinstance(secrets, dict):
                p = doc.add_paragraph()
                run = p.add_run(secrets.get("text", ""))
                self.apply_formatting(run, secrets.get("formatting", {}))
            else:
                doc.add_paragraph(str(secrets))

            # Places Section
            doc.add_heading("Places", level=3)
            for place_name in scenario.get("Places", []):
                place = place_items.get(place_name, {"Name": place_name, "Description": "Unknown Place"})
                if isinstance(place["Description"], dict):
                    description_text = place["Description"].get("text", "Unknown Place")
                else:
                    description_text = place["Description"]
                doc.add_paragraph(f"- {place['Name']}: {description_text}")

            # NPCs Section
            doc.add_heading("NPCs", level=3)
            for npc_name in scenario.get("NPCs", []):
                npc = npc_items.get(npc_name, {"Name": npc_name, "Role": "Unknown",
                                                "Description": {"text": "Unknown NPC", "formatting": {}}})
                p = doc.add_paragraph(f"- {npc['Name']} ({npc['Role']}, {npc.get('Faction', 'Unknown')}): ")
                description = npc['Description']
                if isinstance(description, dict):
                    run = p.add_run(description.get("text", ""))
                    self.apply_formatting(run, description.get("formatting", {}))
                else:
                    p.add_run(str(description))

            # Creatures Section
            doc.add_heading("Creatures", level=3)
            for creature_name in scenario.get("Creatures", []):
                creature = creature_items.get(creature_name, {
                    "Name": creature_name,
                    "Stats": {"text": "No Stats", "formatting": {}},
                    "Powers": {"text": "No Powers", "formatting": {}},
                    "Description": {"text": "Unknown Creature", "formatting": {}}
                })
                stats = creature["Stats"]
                if isinstance(stats, dict):
                    stats_text = stats.get("text", "No Stats")
                else:
                    stats_text = stats
                powers = creature.get("Powers", "Unknown")
                if isinstance(powers, dict):
                    powers_text = powers.get("text", "No Powers")
                else:
                    powers_text = powers
                p = doc.add_paragraph(f"- {creature['Name']} ({stats_text}, {powers_text}): ")
                description = creature["Description"]
                if isinstance(description, dict):
                    run = p.add_run(description.get("text", ""))
                    self.apply_formatting(run, description.get("formatting", {}))
                else:
                    p.add_run(str(description))
        doc.save(file_path)
        messagebox.showinfo("Export Successful", f"Scenario exported successfully to:\n{file_path}")

    def apply_formatting(self, run, formatting):
        if formatting.get('bold'):
            run.bold = True
        if formatting.get('italic'):
            run.italic = True
        if formatting.get('underline'):
            run.underline = True

    def normalize_name(self, name):
        return name.lower().replace('_', ' ').strip()

    def build_portrait_mapping(self):
        mapping = {}
        dir_txt_path = os.path.join("assets", "portraits", "dir.txt")
        if not os.path.exists(dir_txt_path):
            print(f"dir.txt not found at {dir_txt_path}")
            return mapping
        with open(dir_txt_path, "r", encoding="cp1252") as f:
            for line in f:
                line = line.strip()
                if not line.lower().endswith(".png"):
                    continue
                tokens = line.split()
                file_name = tokens[-1]
                if file_name.lower() == "dir.txt":
                    continue
                base_name = os.path.splitext(file_name)[0]
                parts = base_name.split("_")
                filtered_parts = []
                for part in parts:
                    if part.lower() == "portrait" or part.isdigit():
                        continue
                    filtered_parts.append(part)
                if filtered_parts:
                    candidate = " ".join(filtered_parts)
                    normalized_candidate = self.normalize_name(candidate)
                    mapping[normalized_candidate] = file_name
        return mapping

    def associate_npc_portraits(self):
        portrait_mapping = self.build_portrait_mapping()
        if not portrait_mapping:
            print("No portrait mapping was built.")
            return
        db_path = ConfigHelper.get("Database", "path", fallback="default_campaign.db")
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT Name, Portrait FROM npcs")
        npc_rows = cursor.fetchall()
        modified = False
        for npc in npc_rows:
            npc_name = npc["Name"].strip()
            normalized_npc = self.normalize_name(npc_name)
            if normalized_npc in portrait_mapping:
                portrait_file = portrait_mapping[normalized_npc]
                if not npc["Portrait"] or npc["Portrait"].strip() == "":
                    new_portrait_path = os.path.join("assets", "portraits", portrait_file)
                    cursor.execute("UPDATE npcs SET Portrait = ? WHERE Name = ?", (new_portrait_path, npc_name))
                    print(f"Associated portrait '{portrait_file}' with NPC '{npc_name}'")
                    modified = True
        if modified:
            conn.commit()
            print("NPC database updated with associated portraits.")
        else:
            print("No NPC records were updated. Either all have portraits or no matches were found.")
        conn.close()

if __name__ == "__main__":
    app = MainWindow()
    app.mainloop()